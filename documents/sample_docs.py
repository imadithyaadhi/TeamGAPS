import os

def create_invoice():
    invoice_text = (
        "Invoice Number: INV-2025-001\n"
        "Date: June 9, 2025\n\n"
        "Bill To:\n"
        "John Doe\n"
        "Company XYZ\n"
        "123 Business Street, City, Country\n\n"
        "Itemized List:\n"
        "1. Web Development Services - $500\n"
        "2. Domain Registration - $20\n"
        "3. Hosting Plan - $100\n\n"
        "Total Amount Due: $620\n"
        "Due Date: June 15, 2025\n"
        "Payment Method: Bank Transfer\n\n"
        "Thank you for your business!\n"
    )
    with open("invoice.txt", "w", encoding="utf-8") as f:
        f.write(invoice_text)

def create_contract():
    import docx
    doc = docx.Document()
    doc.add_heading("Business Contract", 0)
    doc.add_paragraph("Agreement between [Company ABC] and [Client Name].")
    doc.add_paragraph("Effective Date: June 9, 2025")
    doc.add_paragraph("Terms and Conditions:")
    doc.add_paragraph("1. Scope of Work: Company ABC agrees to provide IT consulting services.", style='List Number')
    doc.add_paragraph("2. Payment Terms: The client agrees to pay $5,000 upon project completion.", style='List Number')
    doc.add_paragraph("3. Confidentiality: Both parties agree to maintain confidentiality.", style='List Number')
    doc.add_paragraph("4. Duration: This agreement is valid from June 9, 2025 to June 9, 2026.", style='List Number')
    doc.add_paragraph("\nSignatures:")
    doc.add_paragraph("[Company Representative]  _______________")
    doc.add_paragraph("[Client Representative]  _______________")
    doc.save("contract.docx")

def create_resume():
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()

    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "John Doe", ln=True, align="C")

    pdf.set_font("Arial", "", 12)
    pdf.cell(0, 10, "Software Engineer", ln=True, align="C")
    pdf.cell(0, 10, "Email: johndoe@example.com | Phone: +1234567890", ln=True, align="C")
    pdf.ln(10)

    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 10, "Experience:", ln=True)
    pdf.set_font("Arial", "", 12)
    pdf.cell(0, 10, "- Senior Developer at TechCorp (2023-2025)", ln=True)
    pdf.cell(0, 10, "- Software Engineer at DevSoft (2020-2023)", ln=True)
    pdf.ln(5)

    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 10, "Skills:", ln=True)
    pdf.set_font("Arial", "", 12)
    pdf.cell(0, 10, "- Python, Java, C++", ln=True)
    pdf.cell(0, 10, "- Machine Learning, AI", ln=True)
    pdf.cell(0, 10, "- Cloud Computing", ln=True)

    pdf.output("resume.pdf")

if __name__ == "__main__":
    create_invoice()
    create_contract()
    create_resume()
    print("Sample files created: invoice.txt, contract.docx, and resume.pdf")